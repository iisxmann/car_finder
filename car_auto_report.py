#!/usr/bin/env python3
"""
Car Auto Report - Fully automated car analysis and report generation.

This script provides a one-shot automated workflow:
1. Ask user for an image path
2. Detect if there's a car in the image (YOLOv8)
3. If yes, classify what type of car it is (EfficientNet-B0)
4. Look up matching vehicles in the local database
5. Generate detailed educational info using TinyLlama
6. Automatically generate Word document (.docx) and PDF report
7. Save everything to the output directory

No intermediate steps or manual file selection required — everything runs automatically!

Usage:
    python3 car_auto_report.py

Models used:
    - YOLOv8n for vehicle detection
    - EfficientNet-B0 for vehicle classification
    - TinyLlama 1.1B for generating detailed descriptions
"""

import argparse
import json
import os
import re
import sys
from datetime import datetime
from pathlib import Path
from typing import Optional

import torch
from PIL import Image
from torchvision.models import efficientnet_b0, EfficientNet_B0_Weights
from ultralytics import YOLO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
from reportlab.lib.colors import HexColor
from llama_cpp import Llama


# Vehicle-related class names from YOLOv8 COCO dataset
VEHICLE_CLASSES = {"car", "bus", "truck", "motorcycle"}

# Vehicle-related categories from ImageNet (EfficientNet)
VEHICLE_CATEGORIES = {
    "ambulance", "convertible", "fire engine", "freight car", "garbage truck",
    "golfcart", "jeep", "limousine", "minibus", "minivan", "moving van",
    "passenger car", "police van", "recreational vehicle", "school bus",
    "sports car", "streetcar", "tow truck", "trailer truck", "trolleybus"
}

# Vehicle category keywords for matching against database models
CATEGORY_KEYWORDS = {
    "sports car": ["sport", "racing", "performance", "gt", "coupe", "brz", "gr86", "gt-r", "camaro", "mustang", "rs"],
    "convertible": ["convertible", "roadster", "spider", "mx-5", "miata"],
    "minibus": ["minibus", "van", "shuttle", "bus"],
    "passenger car": ["sedan", "hatchback", "saloon", "3 series", "5 series", "c-class", "e-class", "a4", "golf", "corolla", "camry", "civic", "accord", "focus", "mazda3"],
    "jeep": ["jeep", "suv", "offroad", "crossover", "x5", "gle", "q5", "rav4", "cr-v", "tucson", "sportage", "forester", "outback", "cx-5", "cx-60"],
    "limousine": ["limousine", "limo"],
    "minivan": ["minivan"],
    "trolleybus": ["trolleybus", "tram"],
    "streetcar": ["streetcar"],
    "electric": ["electric", "ev", "i4", "eqe", "e-tron", "id.4", "ioniq", "ev6", "leaf", "bolt", "mach-e"],
}


class CarAutoAnalyzer:
    """
    Fully automated car analysis pipeline combining:
    - Vehicle detection (YOLOv8)
    - Vehicle classification (EfficientNet-B0)
    - Database lookup
    - LLM description generation (TinyLlama)
    - Report generation (Word + PDF)
    """

    def __init__(self):
        self.yolo_model: Optional[YOLO] = None
        self.efficientnet_model = None
        self.efficientnet_weights = None
        self.llm: Optional[Llama] = None
        self.db: dict = {}
        self._load_all_models()

    def _load_all_models(self) -> None:
        """Load all required models and the database."""
        print("\n" + "=" * 60)
        print("INITIALIZING CAR AUTO REPORT SYSTEM")
        print("=" * 60)
        print("\nLoading offline AI models...")

        # Load YOLOv8 for object detection
        yolo_path = Path(__file__).parent / "yolov8n.pt"
        if yolo_path.exists():
            self.yolo_model = YOLO(str(yolo_path))
            print("  ✓ YOLOv8n loaded (vehicle detection)")
        else:
            print("  ⚠ YOLOv8n not found locally, downloading...")
            self.yolo_model = YOLO("yolov8n.pt")
            print("  ✓ YOLOv8n loaded (vehicle detection)")

        # Load EfficientNet-B0 for classification
        self.efficientnet_weights = EfficientNet_B0_Weights.DEFAULT
        self.efficientnet_model = efficientnet_b0(weights=self.efficientnet_weights)
        self.efficientnet_model.eval()
        print("  ✓ EfficientNet-B0 loaded (vehicle classification)")

        # Load TinyLlama for detailed description generation
        model_path = "/opt/models/tinyllama/tinyllama-1.1b-chat-v1.0.Q4_K_M.gguf"
        if not Path(model_path).exists():
            raise FileNotFoundError(f"TinyLlama model not found at: {model_path}")
        self.llm = Llama(
            model_path=model_path,
            n_ctx=2048,
            n_batch=128,
            verbose=False,
        )
        print("  ✓ TinyLlama loaded (description generation)")

        # Load vehicle database
        db_path = Path(__file__).parent / "data" / "vehicles.json"
        if not db_path.exists():
            raise FileNotFoundError(f"Vehicle database not found: {db_path}")
        with open(db_path, "r", encoding="utf-8") as f:
            self.db = json.load(f)
        print(f"  ✓ Vehicle database loaded ({len(self.db)} brands)")

        print("\n✓ All systems ready!\n")
        print("=" * 60 + "\n")

    def detect_cars(self, image_path: Path) -> dict:
        """Use YOLOv8 to detect if there's a car in the image."""
        if self.yolo_model is None:
            raise RuntimeError("YOLO model not loaded")

        results = self.yolo_model(str(image_path), verbose=False)[0]

        detections = []
        has_car = False

        for box in results.boxes:
            cls_id = int(box.cls[0])
            class_name = self.yolo_model.names[cls_id]
            confidence = float(box.conf[0])

            if class_name in VEHICLE_CLASSES:
                has_car = True
                detections.append({
                    "type": class_name,
                    "confidence": confidence,
                    "bbox": [int(x) for x in box.xyxy[0].tolist()]
                })

        return {
            "has_car": has_car,
            "detections": detections,
            "total_objects_detected": len(results.boxes)
        }

    def classify_car(self, image_path: Path, top_k: int = 5) -> dict:
        """Use EfficientNet-B0 to classify what type of car is in the image."""
        if self.efficientnet_model is None or self.efficientnet_weights is None:
            raise RuntimeError("EfficientNet model not loaded")

        img = Image.open(image_path).convert("RGB")
        preprocess = self.efficientnet_weights.transforms()
        input_tensor = preprocess(img).unsqueeze(0)

        with torch.no_grad():
            output = self.efficientnet_model(input_tensor)

        probabilities = torch.nn.functional.softmax(output[0], dim=0)
        top_prob, top_catid = torch.topk(probabilities, top_k)

        predictions = []
        for i in range(top_prob.size(0)):
            label = self.efficientnet_weights.meta["categories"][top_catid[i]]
            confidence = top_prob[i].item()
            is_vehicle_related = label.lower() in VEHICLE_CATEGORIES

            predictions.append({
                "label": label,
                "confidence": confidence,
                "is_vehicle_related": is_vehicle_related
            })

        return {
            "predictions": predictions,
            "top_vehicle_prediction": next(
                (p for p in predictions if p["is_vehicle_related"]), None
            )
        }

    def find_matching_cars(self, classification: str) -> list[dict]:
        """Find cars in the database that match the given classification."""
        classification_lower = classification.lower()
        matches = []

        keywords = CATEGORY_KEYWORDS.get(classification_lower, [classification_lower])

        for brand, models in self.db.items():
            for model_name, specs in models.items():
                model_lower = model_name.lower()
                brand_lower = brand.lower()

                for kw in keywords:
                    kw_lower = kw.lower()
                    if kw_lower in model_lower or kw_lower in brand_lower:
                        matches.append({
                            "brand": brand,
                            "model": model_name,
                            "engine": specs.get("engine", "N/A"),
                            "hp": specs.get("hp", "N/A"),
                            "fuel": specs.get("fuel", "N/A"),
                            "match_reason": f"Matched keyword '{kw}'"
                        })
                        break

        if len(matches) < 3:
            for brand, models in list(self.db.items())[:5]:
                for model_name, specs in list(models.items())[:2]:
                    entry = {
                        "brand": brand,
                        "model": model_name,
                        "engine": specs.get("engine", "N/A"),
                        "hp": specs.get("hp", "N/A"),
                        "fuel": specs.get("fuel", "N/A"),
                        "match_reason": "Sample vehicle from catalog"
                    }
                    if entry not in matches:
                        matches.append(entry)
                        if len(matches) >= 8:
                            break
                if len(matches) >= 8:
                    break

        return matches[:10]

    def generate_detailed_info(self, classification: str, matching_cars: list[dict]) -> str:
        """Use TinyLlama to generate a detailed, educational description."""
        if self.llm is None:
            raise RuntimeError("LLM not loaded")

        car_list = "\n".join([
            f"- {car['brand']} {car['model']}: {car['engine']}, {car['hp']} hp, {car['fuel']}"
            for car in matching_cars[:6]
        ])

        prompt = f"""You are a helpful automotive expert. A user has an image classified as a "{classification}".

Here are some matching vehicles from our database:
{car_list}

Please provide a detailed, educational response that includes:
1. A brief explanation of what a {classification} is
2. Key characteristics and features of this type of vehicle
3. The main advantages and disadvantages
4. Notable examples or history if relevant

Be informative but concise. Use about 200-300 words."""

        output = self.llm(
            prompt,
            max_tokens=400,
            temperature=0.7,
            stop=["</s>", "User:", "Assistant:"],
        )

        return output["choices"][0]["text"].strip()

    def analyze_image(self, image_path: Path) -> dict:
        """
        Full automated analysis pipeline:
        1. Detect cars
        2. Classify (if car found)
        3. Look up matching vehicles
        4. Generate detailed description
        """
        result = {
            "image_path": str(image_path),
            "image_name": image_path.name,
            "timestamp": datetime.now().isoformat(),
            "detection": None,
            "classification": None,
            "matching_vehicles": [],
            "detailed_description": None,
            "summary": ""
        }

        print(f"Analyzing image: {image_path.name}")
        print("-" * 40)

        # Step 1: Detect if there's a car
        print("Step 1: Detecting vehicles in image...")
        detection_result = self.detect_cars(image_path)
        result["detection"] = detection_result

        if not detection_result["has_car"]:
            result["summary"] = (
                f"No vehicle detected in '{image_path.name}'. "
                f"The image contains {detection_result['total_objects_detected']} object(s), "
                f"none of which are recognized as cars, buses, trucks, or motorcycles."
            )
            print("  ✗ No vehicle detected.")
            return result

        print(f"  ✓ Detected {len(detection_result['detections'])} vehicle(s):")
        for det in detection_result['detections']:
            print(f"      - {det['type'].upper()} ({det['confidence']:.1%})")

        # Step 2: Classify the car type
        print("\nStep 2: Classifying vehicle type...")
        classification_result = self.classify_car(image_path)
        result["classification"] = classification_result

        top_vehicle = classification_result.get("top_vehicle_prediction")
        if top_vehicle:
            print(f"  ✓ Classified as: {top_vehicle['label']} ({top_vehicle['confidence']:.1%})")
        else:
            print("  ⚠ Could not classify as a specific vehicle type.")
            return result

        # Step 3: Find matching vehicles in database
        print("\nStep 3: Looking up matching vehicles in database...")
        classification = top_vehicle['label']
        matching_cars = self.find_matching_cars(classification)
        result["matching_vehicles"] = matching_cars
        print(f"  ✓ Found {len(matching_cars)} matching vehicles")

        # Step 4: Generate detailed description
        print("\nStep 4: Generating detailed description using TinyLlama...")
        detailed_desc = self.generate_detailed_info(classification, matching_cars)
        result["detailed_description"] = detailed_desc
        print("  ✓ Detailed description generated")

        # Build summary
        vehicle_types = [d["type"] for d in detection_result["detections"]]
        result["summary"] = (
            f"Found {len(vehicle_types)} vehicle(s) in '{image_path.name}': {', '.join(set(vehicle_types))}. "
            f"Classified as {classification} ({top_vehicle['confidence']:.1%})."
        )

        return result

    def generate_reports(self, analysis: dict, image_path: Path, output_dir: Path) -> tuple[Path, Path]:
        """Generate both Word and PDF reports from the analysis results."""
        car_name = analysis.get("classification", {}).get("top_vehicle_prediction", {}).get("label", "Unknown Car")
        description = analysis.get("detailed_description", "")
        matching_vehicles = analysis.get("matching_vehicles", [])

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_name = re.sub(r'[^\w\-]', '_', car_name.lower())

        word_path = output_dir / f"{safe_name}_auto_report_{timestamp}.docx"
        pdf_path = output_dir / f"{safe_name}_auto_report_{timestamp}.pdf"

        print("\nStep 5: Generating reports...")
        print("-" * 40)

        # Generate Word document
        self._create_word_report(word_path, car_name, image_path, description, matching_vehicles)
        print(f"  ✓ Word document: {word_path.name}")

        # Generate PDF
        self._create_pdf_report(pdf_path, car_name, image_path, description, matching_vehicles)
        print(f"  ✓ PDF document: {pdf_path.name}")

        return word_path, pdf_path

    def _create_word_report(self, output_path: Path, car_name: str, image_path: Path, description: str, matching_vehicles: list[dict]) -> None:
        """Create a Word document with the report."""
        doc = Document()

        # Title
        title = doc.add_paragraph()
        title_run = title.add_run(car_name.upper())
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 51, 102)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        # Image
        if image_path and image_path.exists():
            try:
                img_para = doc.add_paragraph()
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                with Image.open(image_path) as img:
                    width, height = img.size
                    max_width = 6 * 72
                    if width > max_width:
                        ratio = max_width / width
                        width, height = max_width, height * ratio
                run = img_para.add_run()
                run.add_picture(str(image_path), width=Inches(6))
                doc.add_paragraph()
            except Exception:
                pass

        # Detailed Information
        doc.add_paragraph("DETAILED INFORMATION", style='Heading 1')
        doc.add_paragraph()

        if description:
            doc.add_paragraph(description)
        else:
            doc.add_paragraph("No detailed description available.")

        doc.add_paragraph()

        # Matching Vehicles
        if matching_vehicles:
            doc.add_paragraph("MATCHING VEHICLES FROM CATALOG", style='Heading 1')
            doc.add_paragraph()
            for vehicle in matching_vehicles:
                para = doc.add_paragraph()
                para.add_run(f"• {vehicle['brand']} {vehicle['model']}").bold = True
                para.add_run(f"\n    Engine: {vehicle['engine']}")
                para.add_run(f"\n    Power:  {vehicle['hp']} hp")
                para.add_run(f"\n    Fuel:   {vehicle['fuel']}")
                doc.add_paragraph()

        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer_run = footer.add_run(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | Car Auto Report")
        footer_run.font.size = Pt(9)
        footer_run.font.italic = True
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.save(str(output_path))

    def _create_pdf_report(self, output_path: Path, car_name: str, image_path: Path, description: str, matching_vehicles: list[dict]) -> None:
        """Create a PDF document with the report using reportlab."""
        from reportlab.platypus import Image as RLImage

        styles = getSampleStyleSheet()

        title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=24, spaceAfter=20, alignment=1, textColor=HexColor('#003366'), fontName='Helvetica-Bold')
        heading_style = ParagraphStyle('Heading', parent=styles['Heading2'], fontSize=14, spaceBefore=15, spaceAfter=10, textColor=HexColor('#003366'), fontName='Helvetica-Bold')
        normal_style = ParagraphStyle('Normal', parent=styles['Normal'], fontSize=11, leading=14, spaceAfter=8)
        vehicle_style = ParagraphStyle('Vehicle', parent=styles['Normal'], fontSize=10, leading=13, leftIndent=20)

        story = []

        # Title
        story.append(Paragraph(car_name.upper(), title_style))
        story.append(Spacer(1, 20))

        # Image
        if image_path and image_path.exists():
            try:
                with Image.open(image_path) as img:
                    width, height = img.size
                    max_width = 5.5 * inch
                    if width > max_width:
                        ratio = max_width / width
                        width, height = max_width, height * ratio
                story.append(RLImage(str(image_path), width=width, height=height))
                story.append(Spacer(1, 20))
            except Exception:
                pass

        # Detailed Information
        story.append(Paragraph("DETAILED INFORMATION", heading_style))
        story.append(Spacer(1, 8))

        if description:
            for para in description.split('\n\n'):
                if para.strip():
                    story.append(Paragraph(para.strip(), normal_style))
        else:
            story.append(Paragraph("No detailed description available.", normal_style))

        story.append(Spacer(1, 15))

        # Matching Vehicles
        if matching_vehicles:
            story.append(Paragraph("MATCHING VEHICLES FROM CATALOG", heading_style))
            story.append(Spacer(1, 8))
            for vehicle in matching_vehicles:
                vehicle_text = f"<b>• {vehicle['brand']} {vehicle['model']}</b><br/>    Engine: {vehicle['engine']}<br/>    Power:  {vehicle['hp']} hp<br/>    Fuel:   {vehicle['fuel']}"
                story.append(Paragraph(vehicle_text, vehicle_style))
                story.append(Spacer(1, 8))

        # Footer
        story.append(Spacer(1, 30))
        footer_text = f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | Car Auto Report"
        story.append(Paragraph(footer_text, ParagraphStyle('Footer', fontSize=9, alignment=1, textColor=HexColor('#666666'), fontName='Helvetica-Oblique')))

        doc = SimpleDocTemplate(str(output_path), pagesize=letter, leftMargin=0.75*inch, rightMargin=0.75*inch, topMargin=0.75*inch, bottomMargin=0.75*inch)
        doc.build(story)


def main():
    parser = argparse.ArgumentParser(
        description="Car Auto Report - Fully automated car analysis and report generation",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
This is a one-shot automated tool that handles the entire pipeline:
  1. Image analysis (detect + classify vehicle)
  2. Database lookup for matching vehicles
  3. LLM-powered detailed description generation
  4. Automatic Word and PDF report creation

Just provide an image and let it do all the work!

Example:
  python3 car_auto_report.py
        """
    )
    parser.add_argument(
        "--image", "-i",
        type=Path,
        help="Path to the image to analyze (if not provided, will prompt interactively)"
    )
    parser.add_argument(
        "--output-dir", "-o",
        type=Path,
        default=Path(__file__).parent / "reports",
        help="Directory to save generated reports (default: ./reports/)"
    )
    args = parser.parse_args()

    # Ensure output directory exists
    args.output_dir.mkdir(exist_ok=True)

    # Get image path
    if args.image:
        image_path = args.image
    else:
        print("\nCAR AUTO REPORT - Fully Automated Car Analysis")
        print("=" * 50)
        image_path_str = input("\nEnter the path to an image file: ").strip()
        image_path = Path(image_path_str)

    if not image_path.exists():
        print(f"\nError: Image file not found: {image_path}", file=sys.stderr)
        sys.exit(1)

    # Run the full automated pipeline
    try:
        analyzer = CarAutoAnalyzer()

        # Analyze the image
        analysis = analyzer.analyze_image(image_path)

        if not analysis["detection"]["has_car"]:
            print("\n" + "=" * 60)
            print("RESULT: No vehicle detected in the image.")
            print("=" * 60)
            print(f"\n{analysis['summary']}")
            print("\nPlease try with an image that contains a vehicle.")
            return

        # Generate reports
        word_path, pdf_path = analyzer.generate_reports(analysis, image_path, args.output_dir)

        # Final summary
        print("\n" + "=" * 60)
        print("AUTOMATED ANALYSIS COMPLETE")
        print("=" * 60)
        print(f"\nCar identified: {analysis['classification']['top_vehicle_prediction']['label']}")
        print(f"Confidence: {analysis['classification']['top_vehicle_prediction']['confidence']:.1%}")
        print(f"\nReports generated:")
        print(f"  • {word_path}")
        print(f"  • {pdf_path}")
        print("\nThank you for using Car Auto Report!")

    except Exception as e:
        print(f"\nError during analysis: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
